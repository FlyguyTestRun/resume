import ollama
from docx import Document
from fpdf import FPDF

# Claude generates content (simulate with Ollama for local)
prompt = "Generate resume section for agentic AI skills."
content = ollama.generate(model='phi3:mini', prompt=prompt)['response']

# Save to ReadMe.md
with open('README.md', 'a') as md_file:
    md_file.write("\n### AI Skills\n" + content)

# Save to Word .docx
doc = Document()
doc.add_heading('AI Skills', level=3)
doc.add_paragraph(content)
doc.save('resume.docx')

# Save to PDF
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)
pdf.cell(200, 10, txt=content, ln=1)
pdf.output("resume.pdf")

print("Files updated: README.md, resume.docx, resume.pdf")